# LLM-Agent Hub
# Copyright (C) 2025  Aleksandr Ladygin
# Licensed under the GNU General Public License v3 or later.
#
# See the LICENSE file in the project root for full license information.

from __future__ import annotations

"""
RAG ingestion layer.

Responsibilities:
- read documents from data/
- split into chunks
- embed with the local embedder
- store vectors in Chroma (kb_docs)
- expose ingest_all(), embed_and_store(path), rag_search(), rag_count()

Important: no import-time side effects. Only functions; main.py decides when to call them.
"""

import json
import time
import hashlib
import traceback
from pathlib import Path
from typing import List, Dict, Any, Tuple
from pypdf import PdfReader
from docx import Document

import chromadb
from chromadb.config import Settings

# Local embedder defined in local_embedder.py
from .local_embedder import LocalEmbedder

# --- paths and constants ---
ROOT_DIR   = Path(__file__).resolve().parents[1]  # one level above cognition/
DATA_DIR   = ROOT_DIR / "data"
CHROMA_DIR = ROOT_DIR / ".chroma"
MANIFEST   = ROOT_DIR / "ingest_manifest.json"
COLLECTION = "kb_docs"

DATA_DIR.mkdir(parents=True, exist_ok=True)
CHROMA_DIR.mkdir(parents=True, exist_ok=True)

# Which files are indexed
ALLOWED_EXTS = {".txt", ".md", ".docx", ".pdf", ".json"}

# Files that are never indexed
EXCLUDE_FILES = {
    "insight_log.json",   # reflective notes
    "health_status.json", # core status is not indexed
    "usage_log.jsonl",
    "bridge.log",
}

# create embedder once
_EMBEDDER = LocalEmbedder()  # must expose embed(texts: list[str]) -> list[list[float]]


def _debug(msg: str):
    print(f"[RAG] {msg}", flush=True)


# --- Chroma collection helpers ---
def _get_collection():
    """
    Return the Chroma collection kb_docs.
    Note: embeddings are provided manually, so embedding_function is not configured here.
    """
    client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=Settings(anonymized_telemetry=False),
    )

    coll = client.get_or_create_collection(
        name=COLLECTION,
        metadata={"hnsw:space": "cosine"},
    )
    return coll


# --- manifest helpers for already indexed content ---
def _load_manifest() -> Dict[str, Dict[str, str]]:
    try:
        with open(MANIFEST, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def _save_manifest(m: Dict[str, Dict[str, str]]):
    with open(MANIFEST, "w", encoding="utf-8") as f:
        json.dump(m, f, ensure_ascii=False, indent=2)


# --- file readers ---
def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except Exception:
        _debug("⚠ python-docx dont install. pip install python-docx")
        return ""
    try:
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        _debug(f"❌ DOCX read error {path.name}: {e}")
        return ""
    
def _read_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
        return "\n".join(parts).strip()
    except Exception as e:
        _debug(f"❌ PDF read error {path.name}: {e}")
        return ""

def read_pdf_text(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            # page.extract_text() can return None for some pages
            txt = page.extract_text() or ""
            parts.append(txt)
        return "\n".join(parts).strip()
    except Exception as e:
        return f""  # return empty on error to avoid crashing

def read_docx_text(path: Path) -> str:
    try:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception:
        return ""

def _is_newline_json(text: str) -> bool:
    """Return True if every non-empty line is a valid JSON object."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return False
    for line in lines:
        try:
            json.loads(line)
        except Exception:
            return False
    return True


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".json"}:
        return _read_txt(path).strip()
    if ext == ".docx":
        return _read_docx(path).strip()
    if ext == ".pdf":
        return _read_pdf(path).strip()
    return ""

# --- file hash as version ---
def _file_version_id(path: Path) -> str:
    """
    File content hash. If the file changes, the hash changes.
    Used to decide whether re-indexing is needed.
    """
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()[:16]


# --- text chunker ---
def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
    out = []
    i = 0
    n = len(text)
    while i < n:
        chunk = text[i:i + chunk_size].strip()
        if chunk:
            out.append(chunk)
        i += (chunk_size - overlap)
    return out


# --- main indexing for a single file ---
def embed_and_store(path: Path):
    """
    1. read the file
    2. split into chunks
    3. compute embeddings locally
    4. add chunks to Chroma
    5. record the file version in the manifest
    """
    if path.name in EXCLUDE_FILES:
        _debug(f"[SKIP] {path.name}: excluded by policy.")
        return

    text = extract_text(path)
    if text and _is_newline_json(text):
        return
    if not text:
        _debug(f"[SKIP] {path.name}: empty or unsupported format.")
        return

    chunks = _chunk_text(text)
    if not chunks:
        _debug(f"[SKIP] {path.name}: no chunks produced.")
        return

    # compute embeddings locally
    embeddings = _EMBEDDER.embed(chunks)  # -> list[list[float]]

    fid = _file_version_id(path)

    ids = [f"{fid}_{i}" for i in range(len(chunks))]
    metas = [{"source": path.name} for _ in chunks]

    coll = _get_collection()

    # Before re-indexing the file, remove previous chunks labeled by source
    try:
        coll.delete(where={"source": path.name})
    except Exception:
        # some versions of chroma may not support delete(where=...), ignore failures
        pass

    # add new chunks
    coll.add(
        ids=ids,
        documents=chunks,
        metadatas=metas,
        embeddings=embeddings,
    )

    # persist file version
    manifest = _load_manifest()
    manifest[str(path)] = {
        "fid": fid,
        "ts": str(time.time()),
    }
    _save_manifest(manifest)

    _debug(f"✅ {path.name}: added {len(chunks)} chunks.")


def ingest_all():
    """
    Full indexing of the data/ folder.
    Adds or updates only what changed.
    """
    _debug("ingest_all() started...")
    manifest = _load_manifest()

    for p in DATA_DIR.glob("*"):
        if not p.is_file():
           continue
        if p.name in EXCLUDE_FILES:
           continue
        if p.suffix.lower() not in [".txt", ".docx", ".pdf", ".md"]:
           continue

        fid_now = _file_version_id(p)
        fid_prev = manifest.get(str(p), {}).get("fid")

        if fid_now == fid_prev:
            # unchanged file — skip
            continue

        try:
            embed_and_store(p)
        except Exception as e:
            _debug(f"❌ indexing error {p.name}: {e}")
            traceback.print_exc()

    _debug("ingest_all() finished.")


def rag_search(query: str, k: int = 4) -> List[Dict[str, Any]]:
    """
    Retrieve relevant chunks from Chroma by semantic similarity.
    The query is embedded with the same embedder and compared against Chroma.
    Returns [{'text': ..., 'source': ...}, ...]
    """
    coll = _get_collection()

    q_embed = _EMBEDDER.embed([query])[0]  # query embedding

    # new versions of chroma support query with query_embeddings
    res = coll.query(
        query_embeddings=[q_embed],
        n_results=max(1, k),
    )

    docs  = res.get("documents", [[]])[0]
    metas = res.get("metadatas", [[]])[0]

    out: List[Dict[str, Any]] = []
    for d, m in zip(docs, metas):
        out.append({
            "text": d or "",
            "source": (m or {}).get("source", ""),
        })
    return out


def rag_count() -> int:
    """
    Return the number of elements in the collection (for /status).
    """
    coll = _get_collection()
    try:
        return coll.count()
    except Exception:
        return 0
